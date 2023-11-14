import whisper
from datetime import timedelta
from docx import Document
import streamlit as st

# Выбираем размер модели
model_name = 'base'

@st.cache(allow_output_mutation=True)
def load_model():
    model = whisper.load_model(model_name)

# Выбираем файл
def load_audio():
    uploaded_file = st.file_uploader(label='Выберите аудио для распознавания')
    if uploaded_file is not None:
        audio_data = uploaded_file.getvalue()
        return st.audio(audio_data)
    else:
        return None

def Trasncrib(audio):
    # print(source_file_name)
    # Транскрибация аудио в текст
    result = load_model().transcribe(audio)
    # print(result["text"])
    # Разбиваем полученный текст по сегментам
    segments = result['segments']
    text_massive = []
    for segment in segments:
        startTime = str(0)+str(timedelta(seconds=int(segment['start'])))
        endTime = str(0)+str(timedelta(seconds=int(segment['end'])))
        text = segment['text']
        segmentId = segment['id']+1
        segment = f"{segmentId}. {startTime} - {endTime}\n{text[1:] if text[0] == ' ' else text}"
        st.write(segment)
        text_massive.append(segment)

    print()
    print('Finished')

    # Cохраняем текст с таймингом
    return text_massive

def into_docx(text):
    # Cоздаем новый документ
    doc = Document()
    # Lобавляем параграф с текстом
    doc.add_paragraph(audio + '_' + model_name)
    for key in text:
        doc.add_paragraph(key)
    # Cохраняем документ
    doc.save(audio + '_text_timing_' + model_name + '.docx')


st.title('Транскрибация аудио')
audio = load_audio()
action = st.button('Распознать аудио')
if action:
    Trasncrib(audio)
